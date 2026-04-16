"""
belgeolusturucu.py — GOST-Compliant Word Document Generator
Generates .docx files conforming to Russian GOST R 7.0.97-2016 standards.

Supported Languages: Türkçe (TR), English (EN), Русский (RU)

Features:
  - A4 portrait, configurable margins (archive mode: left=3cm)
  - Times New Roman, selectable font size (12/14/16 pt), black color
  - Justified text, 1.25cm first-line indent
  - Selectable line spacing (1.0 / 1.2 / 1.5)
  - Centered page numbers (Arabic, bottom), hidden on title page
  - Markdown-like heading syntax (# Heading 1, ## Heading 2, ### Heading 3)
  - Centered headings, no trailing period, bold, no indent
  - Native Word XML Multilevel Lists (1. / 1.1. / 1.1.1.) via numbering.xml
  - Bulleted lists (- item / • item) and letter lists (a) item)
  - Title page support (--- TITLE --- block)

Unit Conversions Used:
  1 cm  = 567  twips   (1 inch = 1440 twips, 1 inch = 2.54 cm)
  1 pt  = 20   half-points (for w:sz)

GOST Indent Spec:
  Level 0 (1.)      → number starts at 1.25 cm  → left=1.25cm, hanging=0.75cm → text at 1.25cm
  Level 1 (1.1.)    → number starts at 2.50 cm  → left=2.50cm, hanging=1.00cm → text at 2.50cm
  Level 2 (1.1.1.)  → number starts at 3.75 cm  → left=3.75cm, hanging=1.25cm → text at 3.75cm

  "left" in Word XML = position where text wraps on line 2+.
  "hanging" = how far left the number hangs from "left".
  So number start position = left − hanging.

  Level 0: left = 1.25cm (709tw), hanging = 0.75cm (425tw) → number at 0.50cm offset
           Actually for GOST: number at margin (0cm indent), text at 1.25cm
           → left=709, hanging=709  → number at 0cm, text at 1.25cm ✓

  Level 1: left = 2.50cm (1418tw), hanging = 1.00cm (567tw) → number at 1.50cm
           → left=1418, hanging=709 → number at 1.25cm, text at 2.50cm ✓

  Level 2: left = 3.75cm (2126tw), hanging = 1.00cm (567tw) → number at 2.50cm
           → left=2126, hanging=709 → number at 2.50cm, text at 3.75cm ✓
"""

import re
import sys
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════════════════════
#  LOCALIZATION (title page markers per language)
# ═══════════════════════════════════════════════════════════════

STRINGS = {
    "tr": {
        "title_start": "--- BAŞLIK ---",
        "title_end":   "--- BAŞLIK ---",
    },
    "en": {
        "title_start": "--- TITLE ---",
        "title_end":   "--- TITLE ---",
    },
    "ru": {
        "title_start": "--- ТИТУЛ ---",
        "title_end":   "--- ТИТУЛ ---",
    },
}


# ═══════════════════════════════════════════════════════════════
#  XML HELPERS
# ═══════════════════════════════════════════════════════════════

def add_page_number_field(run):
    """Insert a PAGE field code into the given run (for footer)."""
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    fldChar_text = OxmlElement('w:t')
    fldChar_text.text = "1"

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_text)
    run._r.append(fldChar_end)


def set_run_font(run, font_name="Times New Roman", font_size=14, bold=False):
    """Configure a run's font properties to GOST standards."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    # Ensure font is applied for Cyrillic/Latin via rFonts
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def suppress_page_number_on_first_page(section):
    """Set the section so that page numbering is hidden on the first page (title page)."""
    section.different_first_page_header_footer = True


def set_page_number_start(section, start=1):
    """Set the starting page number for a section."""
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    sectPr.append(pgNumType)


# ═══════════════════════════════════════════════════════════════
#  NATIVE WORD MULTILEVEL LIST SETUP (numbering.xml injection)
# ═══════════════════════════════════════════════════════════════

# We use a single abstractNumId=99 as the template for all GOST multilevel lists.
# Each time we need a fresh list sequence (e.g. "1." restarts), we create a new
# <w:num> element pointing to abstractNumId=99 with level overrides to restart
# counters. This makes Word's "Multilevel List" button active and the hierarchy
# (1. → 1.1. → 1.1.1.) fully native.

ABSTRACT_NUM_ID = 99

def _make_rpr_for_level(font_name, font_size_pt):
    """Create <w:rPr> for a numbering level so auto-generated numbers
    use Times New Roman at the correct size, not the Word default Calibri."""
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_pt * 2))  # half-points
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size_pt * 2))
    rPr.append(szCs)
    return rPr


def setup_multilevel_numbering(doc, font_size=14):
    """Inject the GOST multilevel abstractNum definition into the document's
    numbering.xml. This is called once when building a document.

    GOST R 7.0.97-2016 indent rules (all measured from left page margin):
      Level 0: number at ~0cm,    text body at 1.25cm   → left=709tw,  hanging=709tw
      Level 1: number at ~1.25cm, text body at 2.50cm   → left=1418tw, hanging=709tw
      Level 2: number at ~2.50cm, text body at 3.75cm   → left=2126tw, hanging=709tw
    """
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element

    # Guard: don't inject twice
    for existing in numbering_elem.findall(qn('w:abstractNum')):
        if existing.get(qn('w:abstractNumId')) == str(ABSTRACT_NUM_ID):
            return

    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(ABSTRACT_NUM_ID))

    multiLevelType = OxmlElement('w:multiLevelType')
    multiLevelType.set(qn('w:val'), 'multilevel')
    abstractNum.append(multiLevelType)

    # Level definitions
    # (ilvl, numFmt, lvlText, left_twips, hanging_twips)
    levels = [
        (0, 'decimal', '%1.',       709,  709),   # 1.25cm left, 1.25cm hang → num at 0
        (1, 'decimal', '%1.%2.',    1418, 709),   # 2.50cm left, 1.25cm hang → num at 1.25cm
        (2, 'decimal', '%1.%2.%3.', 2126, 709),   # 3.75cm left, 1.25cm hang → num at 2.50cm
    ]

    for ilvl, num_fmt, text_val, left_tw, hang_tw in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), num_fmt)
        lvl.append(numFmt)

        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), text_val)
        lvl.append(lvlText)

        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # Paragraph properties for this level
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left_tw))
        ind.set(qn('w:hanging'), str(hang_tw))
        pPr.append(ind)
        # Justify text (both = justify in Word)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        lvl.append(pPr)

        # Run properties so auto numbers render in Times New Roman
        lvl.append(_make_rpr_for_level('Times New Roman', font_size))

        abstractNum.append(lvl)

    # Insert abstractNum BEFORE any <w:num> elements (Word spec requirement)
    first_num = numbering_elem.find(qn('w:num'))
    if first_num is not None:
        numbering_elem.insert(list(numbering_elem).index(first_num), abstractNum)
    else:
        numbering_elem.append(abstractNum)


def _get_next_num_id(doc):
    """Find the highest existing numId in numbering.xml and return next available."""
    numbering_elem = doc.part.numbering_part._element
    max_id = 0
    for num_el in numbering_elem.findall(qn('w:num')):
        try:
            nid = int(num_el.get(qn('w:numId')))
            if nid > max_id:
                max_id = nid
        except (TypeError, ValueError):
            pass
    return max(max_id + 1, ABSTRACT_NUM_ID + 1)


def create_new_list_instance(doc, restart_levels=None):
    """Create a new <w:num> element referencing our abstractNum.

    Args:
        doc: The python-docx Document.
        restart_levels: Optional set of ilvl integers to restart at 1.
                       e.g. {0} restarts level 0 counter,
                            {0,1,2} restarts all levels.

    Returns:
        The new numId (int).
    """
    numbering_elem = doc.part.numbering_part._element
    new_id = _get_next_num_id(doc)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_id))

    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), str(ABSTRACT_NUM_ID))
    num.append(abstractNumId)

    # Add level overrides to restart specific levels
    if restart_levels:
        for lvl_idx in sorted(restart_levels):
            lvlOverride = OxmlElement('w:lvlOverride')
            lvlOverride.set(qn('w:ilvl'), str(lvl_idx))
            startOverride = OxmlElement('w:startOverride')
            startOverride.set(qn('w:val'), '1')
            lvlOverride.append(startOverride)
            num.append(lvlOverride)

    numbering_elem.append(num)
    return new_id


def apply_list_numbering(paragraph, num_id, ilvl):
    """Apply native Word XML numbering to a paragraph.

    This sets <w:numPr> with the given numId and ilvl on the paragraph,
    making Word treat it as part of a multilevel list.
    """
    pPr = paragraph._element.get_or_add_pPr()

    # Remove any existing numPr first
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)

    numPr = OxmlElement('w:numPr')

    ilvl_elem = OxmlElement('w:ilvl')
    ilvl_elem.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_elem)

    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(numId_elem)

    pPr.append(numPr)


# ═══════════════════════════════════════════════════════════════
#  LINE PARSERS
# ═══════════════════════════════════════════════════════════════

HEADING_RE = re.compile(r'^(#{1,3})\s+(.+)$')

# Numbered list patterns — most specific first.
# Each pattern captures: group(1) = the number prefix, group(2) = the text content.
#
# Level 3 (ilvl 2): "1.2.3." or "1.2.3 " followed by text
NUMBERED_L3_RE = re.compile(r'^(\d+\.\d+\.\d+)\.?\s+(.+)$')
# Level 2 (ilvl 1): "1.2." or "1.2 " followed by text
NUMBERED_L2_RE = re.compile(r'^(\d+\.\d+)\.?\s+(.+)$')
# Level 1 (ilvl 0): "1." or "1) " followed by text
NUMBERED_L1_RE = re.compile(r'^(\d+)[.)]\s+(.+)$')

LETTER_LIST_RE = re.compile(r'^([a-zA-Zа-яА-ЯёЁ])[.)]\s+(.+)$')
BULLET_LIST_RE = re.compile(r'^[-•–]\s+(.+)$')


def parse_line(line):
    """
    Determine line type and return (type, content, ilvl, number_prefix).

    Types: 'heading', 'numbered', 'letter', 'bullet', 'text', 'empty'

    For 'numbered':
      - content = the text AFTER the number (number stripped since Word generates it)
      - ilvl    = 0, 1, or 2 (Word XML ilvl)
      - number_prefix = the raw numeric prefix like '1', '1.2', '1.2.3'

    For 'letter':
      - content = full original line (letter+text, since we render it manually)
      - number_prefix = the letter character

    For others:
      - number_prefix = ''
    """
    stripped = line.strip()

    if not stripped:
        return ('empty', '', 0, '')

    # Heading: # / ## / ###
    m = HEADING_RE.match(stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2).rstrip('.')  # Remove trailing period per GOST
        return ('heading', text, level, '')

    # Multi-level numbered lists (check deepest level first)
    m = NUMBERED_L3_RE.match(stripped)
    if m:
        return ('numbered', m.group(2), 2, m.group(1))  # ilvl=2

    m = NUMBERED_L2_RE.match(stripped)
    if m:
        return ('numbered', m.group(2), 1, m.group(1))  # ilvl=1

    m = NUMBERED_L1_RE.match(stripped)
    if m:
        return ('numbered', m.group(2), 0, m.group(1))  # ilvl=0

    # Letter list: "a) text" or "б. text"
    m = LETTER_LIST_RE.match(stripped)
    if m:
        return ('letter', stripped, 0, m.group(1))

    # Bulleted list: "- text" or "• text" or "– text"
    m = BULLET_LIST_RE.match(stripped)
    if m:
        return ('bullet', m.group(1), 0, '')

    return ('text', stripped, 0, '')


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════

def build_document(lines, *, archive_mode=False, font_size=14, line_spacing=1.5, lang="tr"):
    """Build and return a python-docx Document from parsed lines."""
    doc = Document()

    # ── Initialize multilevel numbering XML ─────────────────
    setup_multilevel_numbering(doc, font_size=font_size)

    s = STRINGS.get(lang, STRINGS["tr"])

    # ── Section: A4 Portrait with GOST margins ──────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0) if archive_mode else Cm(2.0)
    section.right_margin = Cm(1.0)

    # ── Page numbering starts at 1 ─────────────────────────
    set_page_number_start(section, start=1)

    # ── Normal style: GOST font & paragraph rules ──────────
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(font_size)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)

    pf = style_normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Ensure Times New Roman for all character sets on default style
    rPr = style_normal._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.insert(0, rFonts)

    # ── Detect title page block ─────────────────────────────
    title_lines = []
    body_lines = []
    title_start_tag = s["title_start"].strip().upper()
    title_end_tag = s["title_end"].strip().upper()

    in_title_block = False
    title_closed = False
    for line in lines:
        upper = line.strip().upper()
        if not in_title_block and not title_closed and upper == title_start_tag:
            in_title_block = True
            continue
        if in_title_block and upper == title_end_tag:
            in_title_block = False
            title_closed = True
            continue
        if in_title_block:
            title_lines.append(line)
        else:
            body_lines.append(line)

    # ── Title page (if given) ───────────────────────────────
    has_title_page = len(title_lines) > 0

    if has_title_page:
        # Hide page number on title page
        suppress_page_number_on_first_page(section)

        for i, tl in enumerate(title_lines):
            stripped = tl.strip()
            if not stripped:
                doc.add_paragraph("")  # Preserve blank lines on title page
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(stripped)
            set_run_font(run, font_size=font_size, bold=(i < 3))

        # Add page break after title
        doc.add_page_break()

    # ── Footer: centered page number ────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run()
    set_run_font(footer_run, font_size=font_size)
    add_page_number_field(footer_run)

    # ── State tracking for numbered lists ───────────────────
    # current_num_id:  The active <w:num> numId. None = not in a list.
    # We create a new num instance when:
    #   1. First numbered line after non-list content (current_num_id is None)
    #   2. A level-0 line has number_prefix "1" (explicit restart)
    current_num_id = None

    # ── Process body lines ──────────────────────────────────
    for line in body_lines:
        line_type, content, ilvl, num_prefix = parse_line(line)

        if line_type == 'empty':
            continue

        elif line_type == 'heading':
            # Headings break list context
            current_num_id = None

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(content)
            if ilvl == 1:  # heading level stored in ilvl for headings
                set_run_font(run, font_size=font_size, bold=True)
                run.font.all_caps = True
            elif ilvl == 2:
                set_run_font(run, font_size=font_size, bold=True)
            else:
                set_run_font(run, font_size=font_size, bold=True)

        elif line_type == 'numbered':
            # Decide whether to create a new list instance
            need_new_list = False

            if current_num_id is None:
                # First numbered item after non-list content
                need_new_list = True
            elif ilvl == 0 and num_prefix == '1':
                # Explicit restart: user typed "1. ..." at level 0
                need_new_list = True

            if need_new_list:
                # Restart all three levels
                current_num_id = create_new_list_instance(doc, restart_levels={0, 1, 2})

            p = doc.add_paragraph()

            # Clear the Normal style's first-line indent (numbering XML handles indents)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = None
            p.paragraph_format.line_spacing = line_spacing

            # Apply native XML numbering
            apply_list_numbering(p, current_num_id, ilvl)

            run = p.add_run(content)
            set_run_font(run, font_size=font_size)

        elif line_type == 'letter':
            # Letter lists are rendered manually (not native numbering)
            current_num_id = None
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(2.50)
            # Hanging indent via negative first-line
            p.paragraph_format.first_line_indent = Cm(-0.75)
            run = p.add_run(content)
            set_run_font(run, font_size=font_size)

        elif line_type == 'bullet':
            # Bullets rendered manually with em-dash per GOST
            current_num_id = None
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-0.50)
            run = p.add_run(f"– {content}")  # GOST uses em-dash for bullets
            set_run_font(run, font_size=font_size)

        else:  # 'text'
            # Normal paragraphs break list context
            current_num_id = None
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = line_spacing
            run = p.add_run(content)
            set_run_font(run, font_size=font_size)

    return doc